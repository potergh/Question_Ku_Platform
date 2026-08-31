"""阶段 3：公式导出测试（OMML 原生 / 图片降级 / 降级清单头 / KaTeX 渲染器）。"""

import asyncio
import io
from urllib.parse import quote

from docx import Document

from test_blocks_api import _create_practice, _question
from test_render_service import _load_with_blocks
from test_typography import _save_doc
from app.services import docx_export
from app.services.render_service import build_practice_html
import html as _h


def _formula_doc(latex, display=False):
    node = {"type": "displayFormula" if display else "inlineFormula",
            "attrs": {"latex": latex}}
    return {"type": "doc", "schema_version": 1, "content": [
        {"type": "paragraph", "content": [{"type": "text", "text": "题干："}]},
        ({"type": "paragraph", "content": [node]} if display else
         {"type": "paragraph", "content": [node]}),
    ]}


def _png_1x1() -> bytes:
    from PIL import Image
    buf = io.BytesIO()
    Image.new("RGB", (4, 4), "white").save(buf, "PNG")
    return buf.getvalue()


async def test_top_level_display_formula_save_and_render(client, test_db, tmp_path):
    """验收回归：顶层独立公式保存通过（校验白名单），预览与 Word 均渲染。"""
    practice = await _create_practice(client, test_db, tmp_path)
    q = await _question(client, practice)
    doc = {"type": "doc", "schema_version": 1, "content": [
        {"type": "paragraph", "content": [{"type": "text", "text": "结论："}]},
        {"type": "displayFormula", "attrs": {"latex": "E=mc^2"}},
        {"type": "paragraph", "content": [{"type": "text", "text": "完毕"}]},
    ]}
    res = await _save_doc(client, practice, q, doc)
    assert res.status_code == 200, res.text
    # 旧块反向双写：独立公式回迁为独占一行的 $$…$$
    assert any("$$E=mc^2$$" in (b["content"] or "") for b in res.json()["blocks"])
    # 预览 HTML 输出居中 $$ 定界符（KaTeX 自动渲染）
    p = await _load_with_blocks(test_db, practice["id"])
    html = build_practice_html(p, practice["id"])
    assert "q-formula" in html and "$$E=mc^2$$" in html
    # Word：OMML 行间公式对象；降级清单为空；顶层节点不再丢失
    data, degraded = docx_export.build_docx(p, practice["id"])
    assert "oMathPara" in Document(io.BytesIO(data)).element.xml
    assert degraded == []


async def test_formula_native_omml_no_degraded(client, test_db, tmp_path):
    """可转换公式保持 OMML 原生（可继续编辑），降级清单为空。"""
    practice = await _create_practice(client, test_db, tmp_path)
    q = await _question(client, practice)
    res = await _save_doc(client, practice, q, _formula_doc("a^2+b^2=c^2"))
    assert res.status_code == 200, res.text
    p = await _load_with_blocks(test_db, practice["id"])
    data, degraded = docx_export.build_docx(p, practice["id"])
    assert "oMath" in Document(io.BytesIO(data)).element.xml
    assert degraded == []


async def test_formula_fallback_image_and_header(client, test_db, tmp_path, monkeypatch):
    """OMML 转换失败 → 降级为图片（alt 保留 LaTeX）+ 导出响应头列出清单。"""
    practice = await _create_practice(client, test_db, tmp_path)
    q = await _question(client, practice)
    latex = "\\unsupportedmacro{x}"
    res = await _save_doc(client, practice, q, _formula_doc(latex))
    assert res.status_code == 200, res.text

    monkeypatch.setattr(docx_export, "_latex_to_omml", lambda tex, display: None)
    monkeypatch.setattr(docx_export._FormulaFallback, "render",
                        lambda self, tex, display: _png_1x1())

    # 直接调用：降级清单回传 + 图片带 LaTeX 替代文字
    p = await _load_with_blocks(test_db, practice["id"])
    data, degraded = docx_export.build_docx(p, practice["id"])
    assert degraded == [latex]
    doc = Document(io.BytesIO(data))
    assert any("descr" in sh._inline.xml and latex in sh._inline.xml
               for sh in doc.inline_shapes)

    # 端点：降级清单经 X-Formula-Degraded 头传递（URL 编码）
    res = await client.get(f"/api/practices/{practice['id']}/export/docx")
    assert res.status_code == 200
    assert quote(latex, safe="") in res.headers["x-formula-degraded"]


async def test_formula_fallback_render_failure_keeps_latex(client, test_db, tmp_path, monkeypatch):
    """图片渲染也失败时退回 LaTeX 原文（不静默丢失内容）。"""
    practice = await _create_practice(client, test_db, tmp_path)
    q = await _question(client, practice)
    latex = "\\brokenmacro"
    res = await _save_doc(client, practice, q, _formula_doc(latex))
    assert res.status_code == 200, res.text

    monkeypatch.setattr(docx_export, "_latex_to_omml", lambda tex, display: None)
    monkeypatch.setattr(docx_export._FormulaFallback, "render",
                        lambda self, tex, display: None)

    p = await _load_with_blocks(test_db, practice["id"])
    data, degraded = docx_export.build_docx(p, practice["id"])
    assert degraded == []
    assert f"${latex}$" in Document(io.BytesIO(data)).element.xml or \
           latex in "".join(para.text for para in Document(io.BytesIO(data)).paragraphs)


def test_formula_fallback_real_katex_render():
    """真实渲染：KaTeX + Playwright 把公式截图为 PNG（本地环境已装 chromium）。"""
    fb = docx_export._FormulaFallback()
    try:
        png = fb.render("\\frac{m}{V}", False)
        assert png and png.startswith(b"\x89PNG")
        assert fb.render("\\frac{m}{V}", False) is png   # 缓存命中（同对象）
        assert fb.render("\\badmacro{", False) is None    # KaTeX 语法错误 → None
    finally:
        fb.close()

# ---------------- 阶段 3 收尾专项覆盖（2026-08-31） ----------------
# 覆盖：方程组 / 矩阵 / 分段函数 / 物理单位与向量 / 多行公式 / 长公式溢出检测 / 独立公式跨页保护

# 复杂公式专项：每项都是合法 KaTeX，但 latex2mathml 未必支持（OMML 失败走图片降级，不静默丢失）。
# 注意：Python 普通字符串，\ 表示单个反斜杠、\\ 表示 LaTeX 行分隔符 \。
SPECIALIZED_FORMULAS = {
    "方程组": "\\begin{cases} x+y=3 \\\\ x-y=1 \\end{cases}",
    "矩阵": "\\begin{pmatrix} a & b \\\\ c & d \\end{pmatrix}",
    "分段函数": "f(x)=\\begin{cases} x^2 & x\\ge 0 \\\\ -x & x<0 \\end{cases}",
    "物理单位": "v=3.0\\times10^8\\ \\mathrm{m/s}",
    "向量": "\\vec{F}=m\\vec{a}",
    "多行公式": "\\begin{aligned} a &= b+c \\\\ x &= y-z \\end{aligned}",
}


def _formula_doc_many(formulas, display=True):
    """displayFormula 以顶层块节点组织（真实编辑器结构，doc_render 才会包 .q-formula）。"""
    node_type = "displayFormula" if display else "inlineFormula"
    if display:
        return {"type": "doc", "schema_version": 1, "content": [
            {"type": "paragraph", "content": [{"type": "text", "text": "专项覆盖："}]},
            *[{"type": "displayFormula", "attrs": {"latex": latex}} for latex in formulas],
        ]}
    return {"type": "doc", "schema_version": 1, "content": [
        {"type": "paragraph", "content": [{"type": "text", "text": "专项覆盖："}]},
        *[{"type": node_type, "attrs": {"latex": latex}} for latex in formulas],
    ]}


def _formula_doc_block(latex):
    """独立公式（顶层块节点，真实编辑器结构）。"""
    return {"type": "doc", "schema_version": 1, "content": [
        {"type": "paragraph", "content": [{"type": "text", "text": "题干："}]},
        {"type": "displayFormula", "attrs": {"latex": latex}},
    ]}


async def test_specialized_formula_pipeline_coverage(client, test_db, tmp_path):
    """方程组/矩阵/分段函数/物理单位/向量/多行公式：HTML→Word 全链路不静默丢失。"""
    practice = await _create_practice(client, test_db, tmp_path)
    q = await _question(client, practice)
    latexes = list(SPECIALIZED_FORMULAS.values())
    res = await _save_doc(client, practice, q, _formula_doc_many(latexes))
    assert res.status_code == 200, res.text
    p = await _load_with_blocks(test_db, practice["id"])
    # 预览 HTML：每项 LaTeX 均以 $$…$$ 定界符交给 KaTeX（doc_render 对 & 等做 HTML 转义）
    html = build_practice_html(p, practice["id"])
    for latex in latexes:
        assert f"$${_h.escape(latex)}$$" in html
    # Word：每项要么 OMML 原生、要么进入降级清单、要么保留原文（绝不静默丢）。
    # build_docx 经 to_thread：内部 _FormulaFallback 用 sync Playwright，须在 worker 线程（与生产一致）
    data, degraded = await asyncio.to_thread(docx_export.build_docx, p, practice["id"])
    xml = Document(io.BytesIO(data)).element.xml
    omml_ok = sum(1 for latex in latexes if docx_export._latex_to_omml(latex, True) is not None)
    assert omml_ok == 0 or "oMath" in xml   # 有任意 OMML 路径时文档须含 oMath
    for latex in latexes:
        if docx_export._latex_to_omml(latex, True) is not None:
            continue   # OMML 原生（可继续编辑），无需降级/原文兜底
        # xml 中 & 会被转义为 &amp;，故用 _h.escape 匹配原文兜底
        assert (latex in degraded) or (f"$${_h.escape(latex)}$$" in xml), f"{latex} 丢失"


def test_specialized_formula_katex_render():
    """方程组/矩阵/分段函数/物理单位/向量/多行公式：KaTeX 真实渲染均出图。"""
    fb = docx_export._FormulaFallback()
    try:
        for name, latex in SPECIALIZED_FORMULAS.items():
            png = fb.render(latex, True)
            assert png and png.startswith(b"\x89PNG"), f"{name} KaTeX 渲染失败: {latex}"
    finally:
        fb.close()


async def test_long_formula_overflow_scaled(client, test_db, tmp_path):
    """长公式溢出检测：超出内容宽度的显示公式被自动缩放（zoom），统计 scaled>=1。"""
    from app.services import render_service
    practice = await _create_practice(client, test_db, tmp_path)
    q = await _question(client, practice)
    # 构造明显超出 A4 内容宽度的长公式（连续分式求和）
    long_latex = "S = " + " + ".join(f"\\frac{{{i}^2}}{{{i + 1}}}" for i in range(1, 40))
    res = await _save_doc(client, practice, q, _formula_doc_block(long_latex))
    assert res.status_code == 200, res.text
    p = await _load_with_blocks(test_db, practice["id"])
    html = build_practice_html(p, practice["id"])
    layout = await render_service.render_layout_probe(html, render_service.render_settings(p))
    assert layout["overflow"] >= 1, layout
    assert layout["scaled"] >= 1, layout   # 可缩放范围内已自动适配
    assert layout["unhandled"] == 0, layout


async def test_display_formula_cross_page_protection(client, test_db, tmp_path):
    """独立公式跨页保护：.q-formula / .katex-display 带 break-inside: avoid。"""
    practice = await _create_practice(client, test_db, tmp_path)
    q = await _question(client, practice)
    res = await _save_doc(client, practice, q, _formula_doc_block("E=mc^2"))
    assert res.status_code == 200, res.text
    p = await _load_with_blocks(test_db, practice["id"])
    html = build_practice_html(p, practice["id"])
    assert "page-break-inside: avoid" in html
    assert "break-inside: avoid" in html
    assert ".katex-display" in html

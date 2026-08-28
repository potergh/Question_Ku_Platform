"""Word export service — python-docx generation for student/teacher versions."""

import re
import logging
from pathlib import Path
from io import BytesIO

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

logger = logging.getLogger(__name__)


def generate_word(handout, version: str = "teacher", source_ocr_dirs: dict | None = None) -> BytesIO:
    """
    Generate a Word document from a handout.
    
    Args:
        handout: Handout ORM object with items loaded
        version: "student" (no answers) or "teacher" (with answers)
    
    Returns:
        BytesIO buffer containing the .docx file
    """
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = '微软雅黑'
    font.size = Pt(11)
    
    # Title
    title = doc.add_heading(handout.title, level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Subtitle
    if handout.subject:
        subtitle = doc.add_paragraph(handout.subject)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle.runs[0].font.size = Pt(10)
        subtitle.runs[0].font.color.rgb = RGBColor(128, 128, 128)
    
    doc.add_paragraph()  # Spacing
    
    # Sort items by order
    items = sorted(handout.items, key=lambda i: i.order)
    
    question_num = 1
    for item in items:
        if item.item_type == "section_title":
            _add_section_title(doc, item.custom_content or "")
        
        elif item.item_type in ("question", "example", "exercise"):
            if item.question_snapshot:
                source_id = item.question_snapshot.get("source_id", "")
                ocr_dir = (source_ocr_dirs or {}).get(source_id, "")
                _add_question(doc, item, question_num, version, ocr_dir=ocr_dir)
                question_num += 1
        
        elif item.item_type == "knowledge_note":
            _add_knowledge_note(doc, item.custom_content or "")
    
    # Save to buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def _add_section_title(doc: Document, title: str):
    """Add a section title with styling."""
    heading = doc.add_heading(title, level=2)
    # Add bottom border
    pPr = heading._p.get_or_add_pPr()
    pBdr = pPr.makeelement(qn('w:pBdr'), {})
    bottom = pBdr.makeelement(qn('w:bottom'), {
        qn('w:val'): 'single',
        qn('w:sz'): '8',
        qn('w:space'): '1',
        qn('w:color'): '409EFF',
    })
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_question(doc: Document, item, num: int, version: str, ocr_dir: str = ""):
    """Add a question with optional answer/explanation and images."""
    snap = item.question_snapshot or {}
    
    # Question header
    header_text = f"{num}. "
    if item.item_type == "example":
        header_text = f"【例题】{num}. "
    elif item.item_type == "exercise":
        header_text = f"【练习】{num}. "
    
    # Add question number
    p = doc.add_paragraph()
    run = p.add_run(header_text)
    run.bold = True
    run.font.size = Pt(11)
    
    # Add score if present
    score = snap.get("score")
    if score:
        score_run = p.add_run(f" ({score}分)")
        score_run.font.size = Pt(9)
        score_run.font.color.rgb = RGBColor(128, 128, 128)
    
    # Question content — split text and images
    content = snap.get("content", "") or ""
    _add_content_with_images(doc, content, ocr_dir)
    
    # Options
    options = snap.get("options") or []
    if options:
        for opt in options:
            label = opt.get("label", "")
            opt_content = opt.get("content", "")
            opt_content = _latex_to_text(opt_content)
            doc.add_paragraph(f"{label}. {opt_content}", style='List Bullet')
    
    # Answer section (teacher version only)
    if version == "teacher":
        answer = snap.get("answer", "") or ""
        explanation = snap.get("explanation", "") or ""
        
        if answer or explanation:
            # Add separator
            doc.add_paragraph()
            
            if answer:
                answer = _latex_to_text(answer)
                p = doc.add_paragraph()
                run = p.add_run("【答案】")
                run.bold = True
                run.font.color.rgb = RGBColor(0, 128, 0)
                p.add_run(f" {answer}")
            
            if explanation:
                explanation = _latex_to_text(explanation)
                p = doc.add_paragraph()
                run = p.add_run("【解析】")
                run.bold = True
                run.font.color.rgb = RGBColor(100, 100, 100)
                p.add_run(f" {explanation}")
    else:
        # Student version: add blank space for answer
        doc.add_paragraph("_" * 40)
    
    # Spacing
    doc.add_paragraph()


def _add_knowledge_note(doc: Document, content: str):
    """Add a knowledge note with styling."""
    content = _latex_to_text(content)
    
    # Add note with left border effect (using indentation)
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    
    # Add marker
    run = p.add_run("📝 ")
    run.font.size = Pt(12)
    
    # Add content
    run = p.add_run(content)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(80, 80, 80)
    run.italic = True
    
    doc.add_paragraph()


def _add_content_with_images(doc: Document, content: str, ocr_dir: str = ""):
    """Add content to document, splitting text and images.
    
    Handles asset:// URLs by resolving them to local file paths.
    Converts webp to png for python-docx compatibility.
    """
    if not content:
        return
    
    # Split content into segments: text and images
    # Pattern matches ![alt](url)
    parts = re.split(r'(!\[[^\]]*\]\([^)]+\))', content)
    
    for part in parts:
        img_match = re.match(r'!\[([^\]]*)\]\(([^)]+)\)', part)
        if img_match:
            url = img_match.group(2)
            # Resolve asset:// or /api/ocr-assets/ URLs to local file path
            image_path = _resolve_image_path(url, ocr_dir)
            if image_path and Path(image_path).exists():
                try:
                    # python-docx doesn't support webp — convert to png
                    display_path = image_path
                    if image_path.lower().endswith('.webp'):
                        from PIL import Image
                        import io
                        img = Image.open(image_path)
                        buf = io.BytesIO()
                        img.save(buf, format='PNG')
                        buf.seek(0)
                        display_path = buf
                    
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run()
                    run.add_picture(display_path, width=Inches(4.5))
                except Exception as e:
                    logger.warning(f"Failed to add image {image_path}: {e}")
                    doc.add_paragraph(f"[图片加载失败]")
            else:
                doc.add_paragraph(f"[图片]")
        else:
            # Text segment — convert LaTeX and add
            text = _latex_to_text(part)
            if text.strip():
                doc.add_paragraph(text)


def _resolve_image_path(url: str, ocr_dir: str) -> str | None:
    """Resolve an image URL to a local file path."""
    if not ocr_dir:
        return None
    
    path = None
    if url.startswith("asset://"):
        path = url[len("asset://"):]
        # Normalize double figures/figures/ → figures/
        path = re.sub(r'^figures/figures/', 'figures/', path)
    elif url.startswith("/api/ocr-assets/"):
        parts = url.split("/", 4)  # ['', 'api', 'ocr-assets', 'source_id', 'path']
        if len(parts) >= 5:
            path = parts[4]
            path = re.sub(r'^figures/figures/', 'figures/', path)
    elif url.startswith("file://"):
        return url[7:]  # Strip file:// prefix
    
    if path:
        return str(Path(ocr_dir) / path)
    return None


def _latex_to_text(text: str) -> str:
    """
    Convert LaTeX to plain text for Word.
    Simple conversion: $...$ → content, \frac{a}{b} → a/b, etc.
    Complex formulas remain as text representation.
    """
    if not text:
        return ""
    
    # Remove display math delimiters
    text = re.sub(r'\$\$(.+?)\$\$', r'[\1]', text, flags=re.DOTALL)
    # Remove inline math delimiters
    text = re.sub(r'\$(.+?)\$', r'\1', text)
    
    # Simple LaTeX conversions
    text = text.replace(r'\frac', '')
    text = re.sub(r'\\{([^}]+)\\}{([^}]+)', r'\1/\2', text)  # \frac{a}{b} → a/b
    text = text.replace(r'\times', '×')
    text = text.replace(r'\div', '÷')
    text = text.replace(r'\pm', '±')
    text = text.replace(r'\leq', '≤')
    text = text.replace(r'\geq', '≥')
    text = text.replace(r'\neq', '≠')
    text = text.replace(r'\approx', '≈')
    text = text.replace(r'\infty', '∞')
    text = text.replace(r'\sqrt', '√')
    text = text.replace(r'\pi', 'π')
    text = text.replace(r'\theta', 'θ')
    text = text.replace(r'\alpha', 'α')
    text = text.replace(r'\beta', 'β')
    text = text.replace(r'\gamma', 'γ')
    text = text.replace(r'\Delta', 'Δ')
    text = text.replace(r'\Sigma', 'Σ')
    text = text.replace(r'\Omega', 'Ω')
    
    # Remove remaining backslash commands
    text = re.sub(r'\\[a-zA-Z]+', '', text)
    # Remove curly braces
    text = text.replace('{', '').replace('}', '')
    
    return text.strip()

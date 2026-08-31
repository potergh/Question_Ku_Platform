"""阶段 2：排版体系测试（令牌 / marks 校验 / HTML 与 Word 映射 / 全局默认样式）。"""

import io

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt

from test_blocks_api import _create_practice, _question
from test_render_service import _load_with_blocks
from app.services import docx_export, typography
from app.services.render_service import build_practice_html


# ---------------- 令牌 ----------------

async def test_underline_blank_fill_visible(client, test_db, tmp_path):
    """下划线填空位（纯空格）在预览中不被空白折叠吞掉。"""
    practice = await _create_practice(client, test_db, tmp_path)
    q = await _question(client, practice)
    doc = {"type": "doc", "schema_version": 1, "content": [
        {"type": "paragraph", "content": [
            {"type": "text", "text": "答案是"},
            {"type": "text", "text": "      ", "marks": [{"type": "underline"}]},
            {"type": "text", "text": "。"},
        ]}]}
    res = await _save_doc(client, practice, q, doc)
    assert res.status_code == 200, res.text
    p = await _load_with_blocks(test_db, practice["id"])
    html = build_practice_html(p, practice["id"])
    assert "<u>      </u>" in html
    assert "u { white-space: pre-wrap" in html   # 空格保留规则，下划线才可见


# ---------------- 横线（填写线/分隔线） ----------------

async def test_horizontal_rule_save_and_render(client, test_db, tmp_path):
    """插入横线：校验通过 → HTML 输出 <hr>；旧块降级为一行下划线。"""
    practice = await _create_practice(client, test_db, tmp_path)
    q = await _question(client, practice)
    doc = {"type": "doc", "schema_version": 1, "content": [
        {"type": "paragraph", "content": [{"type": "text", "text": "名前"}]},
        {"type": "horizontalRule"},
        {"type": "paragraph", "content": [{"type": "text", "text": "考号"}]},
    ]}
    res = await _save_doc(client, practice, q, doc)
    assert res.status_code == 200, res.text
    contents = "\n".join(b["content"] or "" for b in res.json()["blocks"])
    assert "＿＿＿＿" in contents              # 旧块降级保留
    p = await _load_with_blocks(test_db, practice["id"])
    html = build_practice_html(p, practice["id"])
    assert '<hr class="q-hr">' in html         # 预览/PDF 输出横线
    assert ".q-hr { border" in html            # 样式定义存在


async def test_horizontal_rule_render_docx(client, test_db, tmp_path):
    """横线在 Word 中渲染为带下边框的段落。"""
    practice = await _create_practice(client, test_db, tmp_path)
    q = await _question(client, practice)
    doc = {"type": "doc", "schema_version": 1, "content": [
        {"type": "paragraph", "content": [{"type": "text", "text": "填空"}]},
        {"type": "horizontalRule"},
    ]}
    res = await _save_doc(client, practice, q, doc)
    assert res.status_code == 200, res.text
    p = await _load_with_blocks(test_db, practice["id"])
    docx = Document(io.BytesIO(docx_export.build_docx(p, practice["id"])[0]))
    has_border = any(x._p.find(qn("w:pPr")) is not None
                     and x._p.find(qn("w:pPr")).find(qn("w:pBdr")) is not None
                     for x in docx.paragraphs)
    assert has_border


def test_size_label_and_font_chain():
    assert typography.size_label(12) == "小四（12 pt）"
    assert typography.size_label(10.5) == "五号（10.5 pt）"
    assert typography.size_label(11) == "11 pt"
    assert typography.css_font_family(None) == '"Times New Roman", serif'
    chain = typography.css_font_family("楷体")
    assert chain.startswith('"Times New Roman"') and '"KaiTi", "楷体"' in chain


def test_practice_default_style_whitelist():
    class Fake:
        page_config = {"default_style": {"font_family": "黑体", "font_size": 12, "line_height": 1.5}}
    assert typography.practice_default_style(Fake()) == {
        "font_family": "黑体", "font_size": 12, "line_height": 1.5}

    class FakeBad:
        page_config = {"default_style": {"font_family": "Comic Sans", "font_size": 13, "line_height": 9}}
    assert typography.practice_default_style(FakeBad()) == typography.DEFAULT_STYLE


# ---------------- 保存校验（排版白名单） ----------------

async def _save_doc(client, practice, q, doc):
    res = await client.put(f"/api/practices/{practice['id']}/questions/{q['id']}/document",
                           json={"document": doc})
    return res


def _styled_doc(**para_attrs):
    return {"type": "doc", "schema_version": 1, "content": [
        {"type": "paragraph", "attrs": para_attrs, "content": [
            {"type": "text", "text": "排版题干", "marks": [
                {"type": "bold"},
                {"type": "textStyle", "attrs": {"fontFamily": "楷体",
                                                 "fontSize": 14, "color": "#f56c6c"}},
            ]},
            {"type": "text", "text": "普通"},
        ]},
        {"type": "bulletList", "content": [
            {"type": "listItem", "content": [
                {"type": "paragraph", "content": [{"type": "text", "text": "列表项一"}]}]},
            {"type": "listItem", "content": [
                {"type": "paragraph", "content": [{"type": "text", "text": "列表项二"}]}]},
        ]},
    ]}


async def test_save_document_accepts_typography(client, test_db, tmp_path):
    practice = await _create_practice(client, test_db, tmp_path)
    q = await _question(client, practice)
    res = await _save_doc(client, practice, q,
                          _styled_doc(textAlign="center", lineHeight=1.5,
                                      spaceBefore=6, firstLineIndent=True))
    assert res.status_code == 200, res.text
    # 列表在旧块方向降级为带前缀的文字行（导出兼容；独立于段落块）
    contents = "\n".join(b["content"] or "" for b in res.json()["blocks"])
    assert "排版题干" in contents
    assert "• 列表项一" in contents and "• 列表项二" in contents


async def test_save_document_rejects_bad_typography(client, test_db, tmp_path):
    practice = await _create_practice(client, test_db, tmp_path)
    q = await _question(client, practice)
    bad_docs = [
        # 字体不在白名单
        {"type": "doc", "schema_version": 1, "content": [{"type": "paragraph", "content": [
            {"type": "text", "text": "x", "marks": [
                {"type": "textStyle", "attrs": {"fontFamily": "Comic Sans"}}]}]}]},
        # 颜色非法
        {"type": "doc", "schema_version": 1, "content": [{"type": "paragraph", "content": [
            {"type": "text", "text": "x", "marks": [
                {"type": "textStyle", "attrs": {"color": "#zzz"}}]}]}]},
        # 字号越界
        {"type": "doc", "schema_version": 1, "content": [{"type": "paragraph", "content": [
            {"type": "text", "text": "x", "marks": [
                {"type": "textStyle", "attrs": {"fontSize": 200}}]}]}]},
        # 段落行距越界
        {"type": "doc", "schema_version": 1, "content": [
            {"type": "paragraph", "attrs": {"lineHeight": 9},
             "content": [{"type": "text", "text": "x"}]}]},
        # 列表项不是 listItem
        {"type": "doc", "schema_version": 1, "content": [
            {"type": "bulletList", "content": [{"type": "paragraph"}]}]},
        # 未知 mark 类型
        {"type": "doc", "schema_version": 1, "content": [{"type": "paragraph", "content": [
            {"type": "text", "text": "x", "marks": [{"type": "blink"}]}]}]},
    ]
    for doc in bad_docs:
        res = await _save_doc(client, practice, q, doc)
        assert res.status_code == 422, doc


# ---------------- marks → HTML（预览/PDF 同源） ----------------

async def test_marks_render_html(client, test_db, tmp_path):
    practice = await _create_practice(client, test_db, tmp_path)
    q = await _question(client, practice)
    res = await _save_doc(client, practice, q,
                          _styled_doc(textAlign="center", lineHeight=1.5, firstLineIndent=True))
    assert res.status_code == 200, res.text
    p = await _load_with_blocks(test_db, practice["id"])
    html = build_practice_html(p, practice["id"])
    assert "KaiTi" in html and "楷体" in html          # 白名单字体链
    assert "font-size:14pt" in html                     # 局部字号
    assert "color:#f56c6c" in html                      # 局部颜色
    assert "<b>" in html                                # 加粗
    assert "text-align:center" in html                  # 段落对齐
    assert "line-height:1.5" in html                    # 行距
    assert "text-indent:2em" in html                    # 首行缩进
    assert '<ul class="q-list">' in html                # 列表
    assert "列表项一" in html


async def test_paragraph_attr_equivalence_defaults(client, test_db, tmp_path):
    """等价默认的属性（如 textAlign=left）不算局部覆盖：渲染输出与无属性一致。"""
    practice = await _create_practice(client, test_db, tmp_path)
    q = await _question(client, practice)
    doc = {"type": "doc", "schema_version": 1, "content": [
        {"type": "paragraph",
         "attrs": {"textAlign": "left", "lineHeight": None, "indent": 0, "firstLineIndent": False},
         "content": [{"type": "text", "text": "默认段落"}]}]}
    res = await _save_doc(client, practice, q, doc)
    assert res.status_code == 200, res.text
    p = await _load_with_blocks(test_db, practice["id"])
    html = build_practice_html(p, practice["id"])
    body = html.split("<body>", 1)[1]                  # 头样式表天然含 text-align，只看题目体
    assert '<div class="q-text" style' not in body      # 等价默认 → 无局部 style 属性
    assert '<div class="q-text"><b>' in body


# ---------------- marks → Word ----------------

async def test_marks_render_docx(client, test_db, tmp_path):
    practice = await _create_practice(client, test_db, tmp_path)
    q = await _question(client, practice)
    res = await _save_doc(client, practice, q, _styled_doc(textAlign="center"))
    assert res.status_code == 200, res.text
    p = await _load_with_blocks(test_db, practice["id"])
    doc = Document(io.BytesIO(docx_export.build_docx(p, practice["id"])[0]))

    para = next(x for x in doc.paragraphs if "排版题干" in x.text)
    assert para.alignment == WD_ALIGN_PARAGRAPH.CENTER          # 段落对齐
    styled = next(r for r in para.runs if r.text == "排版题干")
    assert styled.bold                                           # 加粗
    assert styled.font.size == Pt(14)                            # 局部字号
    assert styled.font.color.rgb is not None
    assert str(styled.font.color.rgb).lower() == "f56c6c"        # 颜色
    assert styled._element.rPr.rFonts.get(qn("w:eastAsia")) == "楷体"   # 东亚字体
    assert any(x.text.startswith("• 列表项") for x in doc.paragraphs)   # 列表降级


# ---------------- 全局默认样式 ----------------

async def test_default_style_applies_html_and_docx(client, test_db, tmp_path):
    practice = await _create_practice(client, test_db, tmp_path)
    pid = practice["id"]
    await client.put(f"/api/practices/{pid}", json={
        "title": practice["title"],
        "page_config": {"default_style": {"font_family": "黑体", "font_size": 12,
                                           "line_height": 2}}})
    p = await _load_with_blocks(test_db, pid)

    html = build_practice_html(p, pid)
    assert '"SimHei"' in html                 # 全局字体进 body
    assert "font-size: 12pt" in html          # 全局字号
    assert "line-height: 2" in html           # 全局行距

    doc = Document(io.BytesIO(docx_export.build_docx(p, pid)[0]))
    normal = doc.styles["Normal"]
    assert normal._element.rPr.rFonts.get(qn("w:eastAsia")) == "黑体"
    assert normal.font.size == Pt(12)
    assert normal.paragraph_format.line_spacing == 2

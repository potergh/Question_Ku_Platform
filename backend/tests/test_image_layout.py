# -*- coding: utf-8 -*-
"""阶段 4：图片排版与编辑测试。

覆盖：layout 数据模型（row/block 默认值）、doc_render 单图/并排/独占/混合、
docx_export 并排表格/独占段落/WebP 转 PNG/缺失资产占位、旧数据迁移补齐 layout。
"""

import io
import json

from PIL import Image

from test_blocks_api import _create_practice, _question, _tiny_webp
from test_render_service import _load_with_blocks
from test_typography import _save_doc
from app.services import docx_export, practice_service
from app.services.doc_render import question_html
from app.services.render_service import build_practice_html
from app.services.rich_document import add_image_layout_default


# ---------------- 构造与断言辅助 ----------------

def _image_doc(specs):
    """specs: [(src, layout, width)] → rich_document。layout/width 为 None 表示不写该属性。"""
    content = [{"type": "paragraph", "content": [{"type": "text", "text": "题干："}]}]
    for src, layout, width in specs:
        attrs = {"src": src, "align": "center"}
        if layout is not None:
            attrs["layout"] = layout
        if width is not None:
            attrs["width"] = width
        content.append({"type": "image", "attrs": attrs})
    return {"type": "doc", "schema_version": 1, "content": content}


def _put_asset(pid, name, data=None):
    """把资产写入练习资产目录，返回 asset://practice/{name}。"""
    if data is None:
        data = _tiny_webp()
    assets = practice_service.practice_assets_dir(pid)
    assets.mkdir(parents=True, exist_ok=True)
    (assets / name).write_bytes(data)
    return "asset://practice/" + name


def _png_bytes():
    buf = io.BytesIO()
    Image.new("RGB", (2, 2)).save(buf, "PNG")
    return buf.getvalue()


async def _asyncio_to_thread(fn, *args):
    import asyncio
    return await asyncio.to_thread(fn, *args)


async def _practice_with_doc(client, test_db, tmp_path, specs):
    """创建练习 → 保存含图文档 → 带块加载。返回 (practice, p)。"""
    practice = await _create_practice(client, test_db, tmp_path)
    q = await _question(client, practice)
    res = await _save_doc(client, practice, q, _image_doc(specs))
    assert res.status_code == 200, res.text
    p = await _load_with_blocks(test_db, practice["id"])
    return practice, p


# ---------------- Task 4.1 / 4.7：layout 模型与迁移 ----------------

def test_add_image_layout_default_is_idempotent():
    doc = {"type": "doc", "schema_version": 1, "content": [
        {"type": "image", "attrs": {"src": "a", "align": "center"}},
        {"type": "image", "attrs": {"src": "b", "align": "center", "layout": "block"}},
        {"type": "paragraph", "content": [
            {"type": "text", "text": "x"},
        ]},
    ]}
    assert add_image_layout_default(doc) == 1          # 仅缺失 layout 的第 1 张补 "row"
    assert doc["content"][0]["attrs"]["layout"] == "row"
    assert doc["content"][1]["attrs"]["layout"] == "block"   # 已有 block 不动
    assert add_image_layout_default(doc) == 0          # 幂等


async def test_migration_adds_layout_via_script_logic(client, test_db, tmp_path):
    """旧数据（image 无 layout）经迁移逻辑补齐为 "row"，已带 block 的不动。"""
    practice = await _create_practice(client, test_db, tmp_path)
    q = await _question(client, practice)
    res = await _save_doc(client, practice, q, _image_doc([
        ("asset://practice/figures/f.webp", None, None),
        ("asset://practice/figures/f.webp", "block", None),
    ]))
    assert res.status_code == 200, res.text

    # 从 DB 读出原始文档（模拟旧数据无 layout）
    from sqlalchemy import select
    from app.models.practice import PracticeQuestion
    async with test_db() as db:
        pq = (await db.execute(select(PracticeQuestion).where(
            PracticeQuestion.practice_id == practice["id"]))).scalars().one()
        doc = json.loads(pq.rich_document)
        # 第 1 张图（content 索引 1）无 layout，第 2 张（索引 2）已是 block
        assert "layout" not in (doc["content"][1].get("attrs") or {})
        assert doc["content"][2]["attrs"]["layout"] == "block"
        n = add_image_layout_default(doc)
        assert n == 1
        assert doc["content"][1]["attrs"]["layout"] == "row"
        assert doc["content"][2]["attrs"]["layout"] == "block"


# ---------------- Task 4.8：doc_render 排版输出 ----------------

async def test_doc_render_single_row_image(client, test_db, tmp_path):
    """单张 row 图：渲染为 .q-img 居中单图（无相邻图不产生行容器）。"""
    practice, p = await _practice_with_doc(client, test_db, tmp_path,
                                           [("asset://practice/f.webp", "row", None)])
    html = build_practice_html(p, practice["id"])
    assert html.count('class="q-img"') >= 1
    assert 'class="q-img-row"' not in html


async def test_doc_render_row_images_side_by_side(client, test_db, tmp_path):
    """连续 3 张 row 图：合并为一个 .q-img-row 行，含 3 个 .q-img-cell 等宽单元格。"""
    practice, p = await _practice_with_doc(client, test_db, tmp_path,
                                           [("asset://practice/f.webp", "row", None)] * 3)
    html = build_practice_html(p, practice["id"])
    assert html.count('class="q-img-row"') == 1
    assert html.count('class="q-img-cell"') == 3


async def test_doc_render_block_image_own_line(client, test_db, tmp_path):
    """layout=block 图：独占一行（.q-img 单图），不并入并排行。"""
    practice, p = await _practice_with_doc(client, test_db, tmp_path,
                                           [("asset://practice/f.webp", "block", None)])
    html = build_practice_html(p, practice["id"])
    assert html.count('class="q-img"') >= 1
    assert 'class="q-img-row"' not in html


async def test_doc_render_mixed_row_block(client, test_db, tmp_path):
    """混合：row,row + block + row → 并排容器、独占图、尾部单张 row 渲染为居中单图。"""
    practice, p = await _practice_with_doc(client, test_db, tmp_path, [
        ("asset://practice/f.webp", "row", None),
        ("asset://practice/f.webp", "row", None),
        ("asset://practice/f.webp", "block", None),
        ("asset://practice/f.webp", "row", None),
    ])
    html = build_practice_html(p, practice["id"])
    assert html.count('class="q-img-row"') == 1      # 仅连续 pair 组成行容器
    assert html.count('class="q-img-cell"') == 2     # 2 张并排
    assert html.count('class="q-img"') == 2          # block 独占 + 尾部单张 row 居中


# ---------------- Task 4.8：docx_export 一致性 ----------------

def _docx_text(xml: str) -> str:
    """粗略抽取 <w:t> 文本。"""
    import re
    return "".join(re.findall(r"<w:t[^>]*>([^<]*)</w:t>", xml))


def _image_rels(doc):
    return [r for r in doc.part.rels.values() if "image" in r.reltype]


async def test_docx_row_and_block_and_webp(client, test_db, tmp_path):
    """Word：连续 row 图进等宽表格行（N 列 1 行），block 图独占段落；WebP 自动转 PNG 可插入。"""
    from docx import Document
    practice = await _create_practice(client, test_db, tmp_path)
    _put_asset(practice["id"], "a.webp", _tiny_webp())
    _put_asset(practice["id"], "b.webp", _tiny_webp())
    _put_asset(practice["id"], "c.webp", _tiny_webp())
    q = await _question(client, practice)
    res = await _save_doc(client, practice, q, _image_doc([
        ("asset://practice/a.webp", "row", None),
        ("asset://practice/b.webp", "row", None),
        ("asset://practice/c.webp", "block", None),
    ]))
    assert res.status_code == 200, res.text
    p = await _load_with_blocks(test_db, practice["id"])
    data, _ = await _asyncio_to_thread(docx_export.build_docx, p, practice["id"])
    doc = Document(io.BytesIO(data))
    # 表格：1 行 2 列（两张 row 并排）；block 图为独立段落
    assert len(doc.tables) >= 1
    assert len(doc.tables[0].columns) == 2
    # 文档含 3 处图片绘制（2 row + 1 block；三图字节相同会被 python-docx 去重为同一 image part，故按 <a:blip> 数断言）
    assert doc.element.xml.count("<a:blip") == 3
    assert "图片缺失" not in doc.element.xml


async def test_docx_missing_asset_placeholder(client, test_db, tmp_path):
    """缺失资产：Word 输出 [图片缺失：name]，不崩溃。"""
    from docx import Document
    practice, p = await _practice_with_doc(client, test_db, tmp_path,
                                           [("asset://practice/ghost.png", "row", None)])
    data, _ = await _asyncio_to_thread(docx_export.build_docx, p, practice["id"])
    xml = Document(io.BytesIO(data)).element.xml
    assert "ghost.png" in _docx_text(xml)
    assert "图片缺失" in xml


async def test_docx_export_png_and_webp_both_supported(client, test_db, tmp_path):
    """PNG 与 WebP 两种格式均可入 Word（PNG 直插、WebP 转 PNG）。"""
    from docx import Document
    practice = await _create_practice(client, test_db, tmp_path)
    _put_asset(practice["id"], "f.webp", _tiny_webp())
    _put_asset(practice["id"], "g.png", _png_bytes())
    q = await _question(client, practice)
    res = await _save_doc(client, practice, q, _image_doc([
        ("asset://practice/f.webp", "row", None),
        ("asset://practice/g.png", "row", None),
    ]))
    assert res.status_code == 200, res.text
    p = await _load_with_blocks(test_db, practice["id"])
    data, _ = await _asyncio_to_thread(docx_export.build_docx, p, practice["id"])
    doc = Document(io.BytesIO(data))
    assert len(_image_rels(doc)) == 2


# ---------------- 阶段 4：并排整行等比缩放（scale 属性） ----------------

def _tall_png_bytes(w=10, h=100):
    buf = io.BytesIO()
    Image.new("RGB", (w, h), (200, 120, 60)).save(buf, "PNG")
    return buf.getvalue()


async def _set_image_scales(test_db, pid, sc=50):
    """给练习文档所有 image 节点补 scale 属性（模拟编辑器整行缩放后保存）。"""
    from sqlalchemy import select
    from app.models.practice import PracticeQuestion
    from app.services.rich_document import serialize
    async with test_db() as db:
        pq = (await db.execute(select(PracticeQuestion).where(
            PracticeQuestion.practice_id == pid))).scalars().one()
        doc = json.loads(pq.rich_document)
        for n in doc["content"]:
            if n.get("type") == "image":
                n["attrs"]["scale"] = sc
        pq.rich_document = serialize(doc)
        await db.commit()


async def test_doc_render_row_scale(client, test_db, tmp_path):
    """row 图片带 scale 属性：行容器内联 width=scale% + margin auto（整行等比缩放并居中）。"""
    practice = await _create_practice(client, test_db, tmp_path)
    _put_asset(practice["id"], "tall.png", _tall_png_bytes())
    q = await _question(client, practice)
    res = await _save_doc(client, practice, q, _image_doc([
        ("asset://practice/tall.png", "row", None),
        ("asset://practice/tall.png", "row", None),
    ]))
    assert res.status_code == 200, res.text
    await _set_image_scales(test_db, practice["id"], 50)
    p = await _load_with_blocks(test_db, practice["id"])
    html = build_practice_html(p, practice["id"])
    assert 'class="q-img-row" style="width:50%;margin:4px auto"' in html
    assert html.count('class="q-img-cell"') == 2


async def test_docx_row_scale(client, test_db, tmp_path):
    """row 图片带 scale 属性：Word 整行等比缩放——行总宽=内容宽*scale%，等宽列随行变窄、表格居中。"""
    from docx import Document
    practice = await _create_practice(client, test_db, tmp_path)
    _put_asset(practice["id"], "tall.png", _tall_png_bytes())
    q = await _question(client, practice)
    res = await _save_doc(client, practice, q, _image_doc([
        ("asset://practice/tall.png", "row", None),
        ("asset://practice/tall.png", "row", None),
    ]))
    assert res.status_code == 200, res.text
    await _set_image_scales(test_db, practice["id"], 50)
    p = await _load_with_blocks(test_db, practice["id"])
    data, _ = await _asyncio_to_thread(docx_export.build_docx, p, practice["id"])
    d = Document(io.BytesIO(data))
    assert d.tables, "no table"
    # 2 图 50%：每列 = 内容宽*50%/2；图片 fit 后不超列宽；表格居中
    from docx.enum.table import WD_TABLE_ALIGNMENT
    col_w = d.tables[0].rows[0].cells[0].width
    widths = [s.width for s in d.inline_shapes]
    assert widths, "no inline pictures"
    assert all(w <= col_w + 1 for w in widths), widths
    assert d.tables[0].alignment == WD_TABLE_ALIGNMENT.CENTER

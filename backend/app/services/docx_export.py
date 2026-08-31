"""Word 导出：练习块 → 可编辑 docx（仅学生版，规格 11.2）。"""

import io
import json
import re
from html.parser import HTMLParser
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.dml.color import RGBColor
from docx.oxml.ns import qn, nsdecls
from docx.oxml import OxmlElement, parse_xml
from docx.shared import Cm, Pt

from app.services import practice_service, typography, workbook_layout
from app.services.doc_render import norm_para_attrs
from app.services.render_service import render_settings, _parse_rich_doc, katex_dist_dir

A4_W, A4_H = Cm(21), Cm(29.7)
MAX_IMG_H = Cm(24)       # 竖长图高度硬封顶（A4 内容区约 24.7cm）
FIT_MAX_H = Cm(8)        # fit 默认上限：与预览同源，避免图片普遍偏大/忽大忽小


# 公式降级渲染页：与预览/PDF 同源 KaTeX，离线 file:// 加载，白底供截图（阶段 3）
_FORMULA_PAGE = (
    '<!doctype html><html><head><meta charset="utf-8">'
    '<link rel="stylesheet" href="{dist}/katex.min.css">'
    '<script src="{dist}/katex.min.js"></script>'
    '<style>html,body{{margin:0;padding:0;background:#fff}}'
    '#f{{display:inline-block;padding:8px}}</style></head>'
    '<body><div id="f"></div></body></html>')


class _FormulaFallback:
    """OMML 转换失败时把公式渲染为图片（懒启动单浏览器会话 + 同公式缓存）。
    渲染失败才退回 LaTeX 原文；降级清单经响应头提示用户（不静默丢失）。"""

    def __init__(self):
        self.degraded: list[str] = []   # 降级为图片的公式（LaTeX 原文）
        self._cache: dict = {}
        self._pw = self._browser = self._page = None
        self._tmp = None
        self._dead = False              # 启动/渲染失败后不再重试（退回原文兼容旧行为）

    def render(self, latex: str, display: bool):
        """返回 PNG 字节；失败返回 None。"""
        import shutil
        import tempfile
        key = (latex, display)
        if key in self._cache:
            return self._cache[key]
        if self._dead:
            return None
        try:
            if self._page is None:
                from playwright.sync_api import sync_playwright
                # 临时目录 + 相对路径加载（set_content 下 file:// 脚本不执行）
                self._tmp = Path(tempfile.mkdtemp())
                shutil.copytree(katex_dist_dir(), self._tmp / "katex")
                (self._tmp / "formula.html").write_text(
                    _FORMULA_PAGE.format(dist="katex", dist2="katex"), encoding="utf-8")
                self._pw = sync_playwright().start()
                self._browser = self._pw.chromium.launch()
                ctx = self._browser.new_context(device_scale_factor=2)
                self._page = ctx.new_page()
                self._page.goto((self._tmp / "formula.html").as_uri(), wait_until="load")
            ok = self._page.evaluate(
                "({latex, display}) => { try {"
                "katex.render(latex, document.getElementById('f'),"
                "{displayMode: display, throwOnError: true}); return true;"
                "} catch (e) { return false; } }", {"latex": latex, "display": display})
            if not ok:
                self._cache[key] = None
                return None
            png = self._page.query_selector("#f").screenshot(type="png")
            self._cache[key] = png
            return png
        except Exception:
            self._dead = True
            self._cache[key] = None
            return None

    def close(self):
        try:
            if self._browser:
                self._browser.close()
        finally:
            if self._pw:
                self._pw.stop()
            if self._tmp:
                import shutil
                shutil.rmtree(self._tmp, ignore_errors=True)


def _insert_formula_image(paragraph, png: bytes, latex: str, display: bool, default_size: float):
    """公式降级图片写入段落：行内随字号定高，独立按宽高比限宽；alt 文字保留 LaTeX 便于追溯。"""
    from PIL import Image
    stream = io.BytesIO(png)
    with Image.open(io.BytesIO(png)) as im:
        px_w, px_h = im.size
    run = paragraph.add_run()
    if display:
        emu_w, emu_h = px_w * 914400 // 192, px_h * 914400 // 192   # 2x 截图 → 96dpi 折算
        width = min(emu_w, int(Cm(15)))
        run.add_picture(stream, width=width)
    else:
        run.add_picture(stream, height=Pt(default_size * 1.5))
    try:   # Word 替代文字：保留原始 LaTeX（图片降级后仍可追溯/手工重建）
        shape = paragraph.runs[-1]._r.find(qn("w:drawing")).find(qn("wp:inline"))
        shape.find(qn("wp:docPr")).set("descr", latex)
    except Exception:
        pass


def build_docx(practice, practice_id: str) -> tuple[bytes, list[str]]:
    """返回 (docx 字节, 降级为图片的公式清单)。"""
    fb = _FormulaFallback()
    try:
        return _build_docx_inner(practice, practice_id, fb), fb.degraded
    finally:
        fb.close()


def _build_docx_inner(practice, practice_id: str, fb: "_FormulaFallback") -> bytes:
    s = render_settings(practice)
    margin = Cm(float(s["margin"].removesuffix("mm")) / 10)   # mm → cm（Cm() 收厘米）
    content_width = A4_W - 2 * margin

    doc = Document()
    # 全局默认样式（阶段 2）：中文白名单字体 + 西文 Times New Roman，设在 Normal 样式上全文生效；
    # 未局部覆盖的内容跟随默认，改全局即生效；局部 marks 在 run 层覆盖。
    ds = s["default_style"]
    normal = doc.styles["Normal"]
    normal.font.name = ds["font_family"] if ds["font_family"] in typography.EN_FONT_CHAIN else typography.DEFAULT_EN_FONT
    normal._element.rPr.rFonts.set(
        qn("w:eastAsia"),
        ds["font_family"] if ds["font_family"] in typography.CN_FONT_CHAIN else typography.DEFAULT_CN_FONT)
    normal.font.size = Pt(ds["font_size"])
    normal.paragraph_format.line_spacing = ds["line_height"]
    sec = doc.sections[0]
    sec.page_width, sec.page_height = A4_W, A4_H
    sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = margin
    if s["show_page_number"]:
        _add_page_number(sec)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(practice.title)
    run.bold = True
    run.font.size = Pt(18)
    _set_cn_font(run)
    if practice.subtitle:
        ps = doc.add_paragraph(practice.subtitle)
        ps.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if s["show_total_score"]:
        total = sum(pq.score or 0 for sec2 in practice.sections for pq in sec2.questions)
        if total > 0:
            pt = doc.add_paragraph(f"满分：{total:g} 分")
            pt.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if s["show_info_bar"]:
        doc.add_paragraph("姓名：____________　班级：____________　日期：____________")

    assets = practice_service.practice_assets_dir(practice_id)
    if practice.layout_document:
        _add_layout_docx(doc, practice, practice_id, assets, content_width, s, fb)
    else:
        for section in practice.sections:
            if section.start_on_new_page:
                bp = doc.add_paragraph()
                bp.add_run().add_break(WD_BREAK.PAGE)
            if section.show_title:
                sp = doc.add_paragraph()
                sr = sp.add_run(section.title)
                sr.bold = True
                sr.font.size = Pt(13)
                _set_cn_font(sr)
            for pq in section.questions:
                rich = _parse_rich_doc(pq)
                if rich is not None:   # 阶段 2：富文档为真源（marks 贯通 Word）
                    _add_question_from_doc(doc, rich, pq, assets, content_width, s, fb)
                else:
                    _add_question(doc, pq, assets, content_width, s, fb)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _add_layout_docx(doc, practice, practice_id, assets: Path, content_width, s: dict, fb):
    """阶段 5：按整册布局块序列导出 Word（与 PDF/预览顺序一致）。"""
    qmap = {pq.id: pq for sec in practice.sections for pq in sec.questions}
    q_no, sub_no = 0, 0
    for blk in practice.layout_document or []:
        t = blk.get("type")
        if t == "subtitle":
            sub_no += 1
            if blk.get("start_on_new_page"):
                bp = doc.add_paragraph()
                bp.add_run().add_break(WD_BREAK.PAGE)
            if blk.get("show_title", True):
                sp = doc.add_paragraph()
                sr = sp.add_run(f"{workbook_layout.section_label(sub_no)}、{blk.get('title') or '小节'}")
                sr.bold = True
                sr.font.size = Pt(13)
                _set_cn_font(sr)
        elif t == "question_ref":
            pq = qmap.get(blk.get("question_id"))
            if pq is None:
                continue
            q_no += 1
            # 用整册序号临时覆盖题号渲染前缀（question_number 在同步时已对齐，此处再保险）
            orig = pq.question_number
            pq.question_number = q_no
            try:
                rich = _parse_rich_doc(pq)
                if rich is not None:
                    _add_question_from_doc(doc, rich, pq, assets, content_width, s, fb)
                else:
                    _add_question(doc, pq, assets, content_width, s, fb)
            finally:
                pq.question_number = orig
        elif t == "custom_text":
            _add_custom_html(doc, blk.get("html") or "", assets, content_width)
        elif t == "spacer":
            h = max(int(blk.get("height") or 20), 4)
            sp = doc.add_paragraph("")
            sp.paragraph_format.space_before = Pt(h * 0.75)
            sp.paragraph_format.space_after = Pt(0)
        elif t == "page_break":
            bp = doc.add_paragraph()
            bp.add_run().add_break(WD_BREAK.PAGE)


def _add_custom_html(doc, html: str, assets: Path, content_width):
    """净化后的自定义内容简单富文本 → Word 段落（p/div/br/strong/em/u/img）。"""
    handler = _CustomHtmlDocx(doc, assets, content_width)
    handler.feed(workbook_layout.sanitize_custom_html(html))
    handler.close()


class _CustomHtmlDocx(HTMLParser):
    """自定义内容 HTML → Word：块级标签起段，行内 strong/em/u 记入 run 样式，img 按内容宽插入。"""
    _BLOCK = {"p", "div", "h1", "h2", "h3", "li", "ol", "ul"}

    def __init__(self, doc, assets: Path, content_width):
        super().__init__(convert_charrefs=True)
        self.doc = doc
        self.assets = assets
        self.content_width = content_width
        self._stack: list = []
        self._runs: list = []
        self._para = None

    def _new_para(self):
        if self._para is None:
            self._para = self.doc.add_paragraph()
        return self._para

    def _flush(self, br: bool = False):
        para = self._new_para()
        for text, b, i, u in self._runs:
            if not text:
                continue
            r = para.add_run(text)
            r.bold = b
            r.italic = i
            r.underline = u
            _set_cn_font(r)
        self._runs = []
        if br:
            para.add_run().add_break()

    def _close_para(self):
        self._flush()
        self._para = None

    def handle_starttag(self, tag, attrs):
        tag = tag.lower()
        if tag in self._BLOCK:
            self._close_para()
            self._new_para()
            if tag.startswith("h") and "b" not in self._stack:
                self._stack.append("b")
        elif tag == "br":
            self._flush(br=True)
        elif tag == "img":
            self._flush()
            src = dict(attrs).get("src", "")
            name = src.removeprefix("asset://practice/")
            if name:
                self._add_img(name)
        elif tag in ("strong", "b"):
            self._stack.append("b")
        elif tag in ("em", "i"):
            self._stack.append("i")
        elif tag == "u":
            self._stack.append("u")
        elif tag in ("span", "s", "sub", "sup"):
            self._stack.append(tag)

    def handle_endtag(self, tag):
        tag = tag.lower()
        if tag in self._BLOCK:
            self._close_para()
            return
        for fmt in ("b", "i", "u", "span", "s", "sub", "sup"):
            if tag == fmt and fmt in self._stack:
                self._stack.remove(fmt)

    def handle_data(self, data):
        if not data:
            return
        self._new_para()
        self._runs.append((data, "b" in self._stack, "i" in self._stack, "u" in self._stack))

    def close(self):
        super().close()
        self._close_para()

    def _add_img(self, name):
        fp = self.assets / name
        if not fp.exists():
            return
        try:
            from PIL import Image
            from docx.shared import Cm
            im = Image.open(fp)
            px_w, px_h = im.size
            max_w = min(int(Cm(8)), int(self.content_width))
            if px_w <= 0 or max_w <= 0:
                return
            ratio = max_w / (px_w * 914400 // 96)
            h = int((px_h * 914400 // 96) * ratio)
            r = self._new_para().add_run()
            with io.BytesIO(fp.read_bytes()) as stream:
                r.add_picture(stream, width=max_w, height=h)
        except Exception:
            pass


def _set_cn_font(run, name="宋体"):
    """中文用指定字体，西文统一 Times New Roman。"""
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)


def _set_run_font(run, name: str):
    """白名单字体 → run：中文字体写东亚属性，西文写 ascii（中英混排各取所需）。"""
    if name in typography.CN_FONT_CHAIN:
        run.font.name = typography.DEFAULT_EN_FONT
        run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    else:
        run.font.name = name


def _add_page_number(section):
    """页脚居中插入 PAGE 域（Word 打开后显示真实页码）。"""
    p = section.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    for tag, attr, val in [("w:fldChar", "begin", None), ("w:instrText", None, "PAGE"),
                           ("w:fldChar", "end", None)]:
        el = OxmlElement(tag)
        if tag == "w:fldChar":
            el.set(qn("w:fldCharType"), attr)
        else:
            el.text = val
        run._r.append(el)


def _add_question(doc, pq, assets: Path, content_width, s: dict, fb):
    prefix = f"{pq.question_number}. "
    if s["show_score"] and pq.score is not None:
        prefix += f"（{pq.score:g} 分）"
    blocks = sorted(pq.blocks, key=lambda b: b.position)
    # 无文字块（纯图片题）：题号单独一行并置顶；否则并入首个文字段（题号与题干同行）
    prefix_used = not any(b.block_type == "text" for b in blocks)
    if prefix_used:
        doc.add_paragraph(prefix)
    imgs: list = []

    def flush_imgs():
        if not imgs:
            return
        if len(imgs) == 1:
            _add_image(doc, imgs[0].content or "", imgs[0].style_config or {}, assets, content_width)
        else:
            _add_image_row(doc, [(b.content or "", b.style_config or {}) for b in imgs],
                           assets, content_width)
        imgs.clear()

    for b in blocks:
        if b.block_type == "image":
            imgs.append(b)
            continue
        flush_imgs()
        style = b.style_config or {}
        if b.block_type == "text":
            tp = doc.add_paragraph()
            # 公式内嵌为 OMML 数学对象（转换失败降级为公式图片，再失败保留 LaTeX 原文）
            _add_rich_runs(tp, ("" if prefix_used else prefix)
                           + (b.content or "").replace("**", ""), assets, fb,
                           s["default_style"]["font_size"])
            prefix_used = True
        elif b.block_type == "options":
            try:
                opts = json.loads(b.content) if b.content else []
            except (TypeError, json.JSONDecodeError):
                opts = []
            for o in opts:
                op = doc.add_paragraph()
                op.paragraph_format.left_indent = Cm(0.74)
                op.add_run(f"{o.get('label', '')}. ")
                _add_rich_runs(op, o.get("content", ""), assets, fb, s["default_style"]["font_size"])
        elif b.block_type == "answer_space":
            for _ in range(int(style.get("rows", 0))):
                doc.add_paragraph("")
        # answer/explanation 块学生版不输出
    flush_imgs()


# ---------------- 阶段 2：rich_document → Word（marks/段落属性贯通） ----------------

_ALIGN_MAP = {"left": WD_ALIGN_PARAGRAPH.LEFT, "center": WD_ALIGN_PARAGRAPH.CENTER,
              "right": WD_ALIGN_PARAGRAPH.RIGHT, "justify": WD_ALIGN_PARAGRAPH.JUSTIFY}


def _add_question_from_doc(doc, rich: dict, pq, assets: Path, content_width, s: dict, fb):
    """从富文档生成 Word 内容；无属性连续段落合并为单段（与旧块粒度一致）。"""
    default_size = s["default_style"]["font_size"]
    prefix = f"{pq.question_number}. "
    if s["show_score"] and pq.score is not None:
        prefix += f"（{pq.score:g} 分）"
    prefix_used = False
    pending: list = []
    imgs: list = []

    def flush_paras():
        nonlocal prefix_used
        if not pending:
            return
        p = doc.add_paragraph()
        if not prefix_used:
            p.add_run(prefix)
            prefix_used = True
        for i, nodes in enumerate(pending):
            if i:
                p.add_run().add_break(WD_BREAK.LINE)
            _add_inline_runs(p, nodes, assets, fb, default_size)
        pending.clear()

    def flush_imgs():
        if not imgs:
            return
        if len(imgs) == 1:
            _add_image(doc, imgs[0][0], imgs[0][1], assets, content_width)
        else:
            _add_image_row(doc, imgs, assets, content_width)
        imgs.clear()

    for node in rich.get("content") or []:
        t = node.get("type")
        if t == "paragraph":
            attrs = norm_para_attrs(node.get("attrs"))
            if not attrs:
                pending.append(node.get("content") or [])
                continue
            flush_paras(); flush_imgs()
            p = doc.add_paragraph()
            if not prefix_used:
                p.add_run(prefix)
                prefix_used = True
            _apply_para_attrs(p, attrs, default_size)
            _add_inline_runs(p, node.get("content") or [], assets, fb, default_size)
        elif t == "image":
            flush_paras()
            attrs = node.get("attrs") or {}
            if attrs.get("layout") == "block":
                flush_imgs()
                _add_image(doc, attrs.get("src", ""), attrs, assets, content_width)
            else:
                imgs.append((attrs.get("src", ""), attrs))
        elif t == "displayFormula":
            flush_paras(); flush_imgs()
            fp = doc.add_paragraph()
            fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _add_formula(fp, (node.get("attrs") or {}).get("latex", ""), True, fb, default_size)
        elif t == "optionGroup":
            flush_paras(); flush_imgs()
            for o in node.get("content") or []:
                if o.get("type") != "option":
                    continue
                op = doc.add_paragraph()
                op.paragraph_format.left_indent = Cm(0.74)
                op.add_run(f"{(o.get('attrs') or {}).get('label', '')}. ")
                _add_inline_runs(op, o.get("content") or [], assets, fb, default_size)
        elif t == "answerSpace":
            flush_paras(); flush_imgs()
            for _ in range(int((node.get("attrs") or {}).get("rows", 0))):
                doc.add_paragraph("")
        elif t == "horizontalRule":
            flush_paras(); flush_imgs()
            _add_horizontal_rule(doc)
        elif t in ("bulletList", "orderedList"):
            flush_paras(); flush_imgs()
            _add_list(doc, node, assets, 0, fb, default_size)
    flush_paras(); flush_imgs()
    if not prefix_used:   # 无文字内容（纯图片题）：题号单独一行
        doc.add_paragraph(prefix)


def _add_horizontal_rule(doc):
    """横线：一个带下边框的段落（与预览/PDF 的 hr 视觉一致）。"""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")      # 1pt 实线（sz 以 1/8 pt 计）
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def _add_list(doc, node: dict, assets: Path, depth: int, fb, default_size: float):
    """列表 → 缩进段落 + 符号/编号前缀（Word 原生列表样式不在首版范围）。"""
    ordered = node.get("type") == "orderedList"
    for i, item in enumerate(node.get("content") or []):
        if item.get("type") != "listItem":
            continue
        first = True
        for inner in item.get("content") or []:
            if inner.get("type") == "paragraph":
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(0.74 * (depth + 1))
                if first:
                    p.add_run(f"{i + 1}. " if ordered else "• ")
                    first = False
                _add_inline_runs(p, inner.get("content") or [], assets, fb, default_size)
            elif inner.get("type") in ("bulletList", "orderedList"):
                _add_list(doc, inner, assets, depth + 1, fb, default_size)


def _add_inline_runs(paragraph, nodes, assets: Path, fb=None, default_size: float = 11):
    """行内节点 → runs：marks → run 属性；公式 → OMML（失败降级图片）；行内图随基线。"""
    for n in nodes or []:
        t = n.get("type")
        if t == "text":
            marks = n.get("marks") or []
            text = n.get("text", "")
            if not any(m.get("type") == "bold" for m in marks):
                text = text.replace("**", "")   # 旧 Markdown 加粗残留剔除（与旧路径一致）
            _apply_marks(paragraph.add_run(text), marks)
        elif t == "hardBreak":
            paragraph.add_run().add_break(WD_BREAK.LINE)
        elif t == "inlineImage":
            name = ((n.get("attrs") or {}).get("src") or "").removeprefix("asset://practice/")
            path = assets / name
            if path.exists():
                paragraph.add_run().add_picture(_picture_source(path), height=Cm(0.9))
            else:
                paragraph.add_run(f"[图缺失:{name}]")
        elif t == "inlineFormula":
            latex = (n.get("attrs") or {}).get("latex", "")
            _add_formula(paragraph, latex, False, fb, default_size)
        elif t == "displayFormula":
            latex = (n.get("attrs") or {}).get("latex", "")
            _add_formula(paragraph, latex, True, fb, default_size)


def _add_formula(paragraph, latex: str, display: bool, fb, default_size: float):
    """公式三级兜底：OMML 原生（可继续编辑）→ 图片降级（保显示，记入降级清单）→ LaTeX 原文。"""
    el = _latex_to_omml(latex, display)
    if el is not None:
        paragraph._p.append(el)
        return
    png = fb.render(latex, display) if fb is not None else None
    if png:
        _insert_formula_image(paragraph, png, latex, display, default_size)
        fb.degraded.append(latex)
    else:
        paragraph.add_run(("$$" if display else "$") + latex + ("$$" if display else "$"))


def _apply_marks(run, marks):
    for m in marks or []:
        t = m.get("type")
        if t == "bold":
            run.bold = True
        elif t == "italic":
            run.italic = True
        elif t == "underline":
            run.underline = True
        elif t == "strike":
            run.font.strike = True
        elif t == "superscript":
            run.font.superscript = True
        elif t == "subscript":
            run.font.subscript = True
        elif t == "textStyle":
            attrs = m.get("attrs") or {}
            if attrs.get("fontFamily"):
                _set_run_font(run, attrs["fontFamily"])
            if isinstance(attrs.get("fontSize"), (int, float)):
                run.font.size = Pt(attrs["fontSize"])
            if attrs.get("color"):
                run.font.color.rgb = RGBColor.from_string(attrs["color"].lstrip("#"))


def _apply_para_attrs(p, attrs: dict, default_size: float):
    if attrs.get("textAlign"):
        p.alignment = _ALIGN_MAP.get(attrs["textAlign"])
    if attrs.get("lineHeight"):
        p.paragraph_format.line_spacing = attrs["lineHeight"]
    if attrs.get("spaceBefore"):
        p.paragraph_format.space_before = Pt(attrs["spaceBefore"])
    if attrs.get("spaceAfter"):
        p.paragraph_format.space_after = Pt(attrs["spaceAfter"])
    if attrs.get("firstLineIndent"):
        p.paragraph_format.first_line_indent = Pt(2 * default_size)   # 首行缩进两字符（随默认字号）
    if attrs.get("indent"):
        p.paragraph_format.left_indent = Cm(0.74 * attrs["indent"])


# 公式段：$$…$$ / \[…\] 为行间，$…$ / \(…\) 为行内（先长后短，正则分支顺序即优先级）
_MATH_RE = re.compile(r"\$\$(.+?)\$\$|\\\[(.+?)\\\]|\\\((.+?)\\\)|\$([^\$\n]+?)\$", re.S)


def _split_math(content: str):
    """内容切分为文字/公式段，返回 (是否公式, 文本, 是否行间)。"""
    segs, last = [], 0
    for m in _MATH_RE.finditer(content):
        if m.start() > last:
            segs.append((False, content[last:m.start()], False))
        if m.group(1) is not None or m.group(2) is not None:
            segs.append((True, m.group(1) or m.group(2), True))
        else:
            segs.append((True, m.group(3) or m.group(4), False))
        last = m.end()
    if last < len(content):
        segs.append((False, content[last:], False))
    return segs


def _latex_to_omml(tex: str, display: bool):
    """LaTeX → OMML 元素（Word 原生公式）；转换失败返回 None，调用方退回原文。
    注：mathml2omml 对个别命令（如 \\vec）会产出标签不闭合的坏 OMML，
    parse_xml 一并纳入 try——坏 OMML 按转换失败处理，走图片降级，避免污染整个 docx。"""
    try:
        import latex2mathml.converter
        import mathml2omml
        omml = mathml2omml.convert(latex2mathml.converter.convert(tex.strip()))
        if display:
            omml = f"<m:oMathPara>{omml}</m:oMathPara>"
        return parse_xml(f'<root {nsdecls("m")}>{omml}</root>')[0]
    except Exception:
        return None


def _add_rich_runs(paragraph, content: str, assets: Path, fb=None, default_size: float = 11):
    """混合富内容：公式 → 内嵌 OMML（失败降级图片），图片引用 → 行内图片，其余纯文字。"""
    for is_math, text, display in _split_math(content or ""):
        if is_math:
            _add_formula(paragraph, text, display, fb, default_size)
            continue
        _add_text_with_imgs(paragraph, text, assets)


def _add_text_with_imgs(paragraph, content: str, assets: Path):
    """文字中的图片引用 → 行内图片（随文字基线排版）。"""
    last = 0
    for m in re.finditer(r"!\[[^\]]*\]\((asset://[^\s\)]+)\)|(asset://[^\s\)]+)", content or ""):
        if m.start() > last:
            paragraph.add_run(content[last:m.start()])
        name = (m.group(1) or m.group(2)).removeprefix("asset://practice/")
        path = assets / name
        if path.exists():
            paragraph.add_run().add_picture(_picture_source(path), height=Cm(0.9))
        else:
            paragraph.add_run(f"[图缺失:{name}]")
        last = m.end()
    if last < len(content or ""):
        paragraph.add_run(content[last:])


def _add_image(doc, src: str, style: dict, assets: Path, content_width):
    name = (src or "").removeprefix("asset://practice/")
    path = assets / name
    ip = doc.add_paragraph()
    align = style.get("align", "center")
    ip.alignment = {"left": WD_ALIGN_PARAGRAPH.LEFT, "right": WD_ALIGN_PARAGRAPH.RIGHT}.get(
        align, WD_ALIGN_PARAGRAPH.CENTER)
    if not path.exists():
        ip.add_run(f"[图片缺失：{name}]")
        return
    width = None
    w = style.get("width", "fit")
    if isinstance(w, (int, float)):
        width = content_width * w / 100
    elif isinstance(w, str) and w.endswith("%"):
        width = content_width * float(w.removesuffix("%")) / 100
    if width is None:   # fit：原尺寸，受默认上限（内容区 50% 宽 / 8cm 高）与硬封顶约束
        width = _fit_width(path, content_width)
    run = ip.add_run()
    run.add_picture(_picture_source(path), width=width)


def _add_image_row(doc, items, assets: Path, content_width):
    """连续图片并排：无边框单行表格，等宽列；单元格内居中且不超过列宽。items 为 (src, style) 序列。
    阶段 4：scale（整行等比缩放）→ 行总宽 = 内容宽 * scale%，等宽列随行变窄，表格保持居中。"""
    n = len(items)
    scale = (items[0][1] or {}).get("scale")
    row_w = content_width
    if isinstance(scale, (int, float)) and 0 < scale < 100:
        row_w = int(content_width * scale / 100)
    table = doc.add_table(rows=1, cols=n)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell_w = max(1, row_w // n)
    for i, (src, _) in enumerate(items):
        cell = table.rows[0].cells[i]
        cell.width = cell_w
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        name = (src or "").removeprefix("asset://practice/")
        path = assets / name
        if not path.exists():
            p.add_run(f"[图片缺失：{name}]")
            continue
        width = _fit_width(path, cell_w)
        p.add_run().add_picture(_picture_source(path), width=width)


def _natural_size(path: Path):
    """图片自然宽高（EMU）。忽略文件内嵌 dpi（各图 96/300 不一导致忽大忽小），统一按 96 折算。"""
    try:
        from PIL import Image
        with Image.open(path) as im:
            px_w, px_h = im.size
        return int(px_w / 96 * 914400), int(px_h / 96 * 914400)
    except Exception:
        return None


def _fit_width(path: Path, max_width, max_height=None):
    """fit：宽受 max_width 与默认上限（内容区 50%）双重封顶，高度封顶后再反推宽度；读不到尺寸保持原样。
    max_height：可选的高度封顶，先按它等比缩宽再走默认上限（阶段 4 行内高度封顶已由 scale 取代，仅保留参数兼容）。"""
    ns = _natural_size(path)
    if not ns:
        return None
    w, h = ns
    cap_w = min(max_width, int(max_width / 2))
    if max_height and h > max_height:
        w = int(w * max_height / h)
        h = max_height
    if h > FIT_MAX_H:
        w = int(w * FIT_MAX_H / h)
    if h > MAX_IMG_H:
        w = min(w, int(w * MAX_IMG_H / h))
    if w >= cap_w:
        return cap_w
    return w if (w != ns[0] or h != ns[1]) else None


def _picture_source(path: Path):
    """docx 只识别常见光栅格式；webp/avif 等（OCR 常见）先用 PIL 转 PNG 流。"""
    if path.suffix.lower() in {".png", ".jpg", ".jpeg", ".gif", ".bmp", ".tif", ".tiff", ".ico"}:
        return str(path)
    from PIL import Image
    buf = io.BytesIO()
    Image.open(path).convert("RGB").save(buf, "PNG")
    buf.seek(0)
    return buf

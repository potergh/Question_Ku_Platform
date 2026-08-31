"""HTML 组装单测：不经浏览器，直接断言 HTML 文本。"""

from docx import Document
import io

from test_blocks_api import _create_practice
from app.services import docx_export
from app.services.render_service import build_practice_html, render_settings


async def _full_practice(client, test_db, tmp_path):
    practice = await _create_practice(client, test_db, tmp_path)
    detail = (await client.get(f"/api/practices/{practice['id']}/detail")).json()
    return detail


async def _load_with_blocks(test_db, pid):
    """test_db 是 session 工厂；直接查测试库，勿用生产 get_db。"""
    from sqlalchemy import select
    from sqlalchemy.orm import selectinload
    from app.models.practice import Practice, PracticeSection, PracticeQuestion
    async with test_db() as db:
        return (await db.execute(
            select(Practice).where(Practice.id == pid)
            .options(selectinload(Practice.sections)
                     .selectinload(PracticeSection.questions)
                     .selectinload(PracticeQuestion.blocks))
            .execution_options(populate_existing=True)
        )).scalar_one()


def test_render_settings_defaults():
    class FakePractice:
        page_config = None
    s = render_settings(FakePractice())
    assert s["margin"] == "25mm"
    assert s["show_info_bar"] is True
    assert s["show_page_number"] is True
    assert s["show_score"] is False


async def test_build_html_contains_structure(client, test_db, tmp_path):
    detail = await _full_practice(client, test_db, tmp_path)
    practice = await _load_with_blocks(test_db, detail["id"])
    html = build_practice_html(practice, detail["id"])
    assert detail["title"] in html
    assert "姓名" in html                      # 默认显示学生信息栏
    assert "1." in html                        # 题号
    assert "q-option" in html                  # 单选题必带选项块
    assert "answer-space" in html              # 留白块
    assert "katex" in html                     # KaTeX 标签


async def test_answer_space_blank_no_lines(client, test_db, tmp_path):
    """答题留白只留空白，不画横线（用户决策 2026-08-30）。"""
    detail = await _full_practice(client, test_db, tmp_path)
    practice = await _load_with_blocks(test_db, detail["id"])
    html = build_practice_html(practice, detail["id"])
    assert "answer-space" in html
    assert "space-line" in html
    assert "border-bottom" not in html


async def test_build_html_respects_page_config(client, test_db, tmp_path):
    detail = await _full_practice(client, test_db, tmp_path)
    pid = detail["id"]
    await client.put(f"/api/practices/{pid}", json={"title": detail["title"],
        "page_config": {"show_info_bar": False, "show_score": True, "margin_preset": "narrow"}})
    practice = await _load_with_blocks(test_db, pid)
    html = build_practice_html(practice, pid)
    settings = render_settings(practice)
    assert "姓名" not in html                  # 信息栏关闭
    assert settings["margin"] == "15mm"


async def test_prefix_inline_and_img_row(client, test_db, tmp_path):
    """题号并入题干首行；连续多图并排一行（HTML 与 Word 一致）。"""
    # 连续两图 → materialize 出两个相邻 image 块（题号 3 顺带验证重排为 1）
    from app.models import Source, Question
    from test_blocks_api import _tiny_webp
    ocr_dir = tmp_path / "ocr" / "d"
    (ocr_dir / "figures").mkdir(parents=True, exist_ok=True)
    (ocr_dir / "figures" / "f.webp").write_bytes(_tiny_webp())
    async with test_db() as db:
        source = Source(filename="t.pdf", file_path="/tmp/t.pdf", file_type="pdf",
                        ocr_status="done", ocr_result_path=str(ocr_dir))
        db.add(source)
        await db.commit()
        q = Question(source_id=source.id, source_question_id="Q1", question_number=3,
                     question_type="short_answer",
                     content="题干首行\n![图](asset://figures/f.webp)![图](asset://figures/f.webp)",
                     options=None)
        db.add(q)
        await db.commit()
        await db.refresh(q)
        qid = q.id
    practice = (await client.post("/api/practices", json={
        "title": "t", "from_basket": False, "question_ids": [qid]})).json()
    pid = practice["id"]
    await client.get(f"/api/practices/{pid}/detail")   # 懒物化 + 重排题号 3→1
    p = await _load_with_blocks(test_db, pid)

    html = build_practice_html(p, pid)
    assert '<div class="q-text"><b>1. </b>题干首行</div>' in html   # 题号与题干同行（无 ![图]( 残留）
    assert "q-img-row" in html                                       # 连续两图并排
    assert html.count('<div class="q-img-cell">') == 2

    doc = Document(io.BytesIO(docx_export.build_docx(p, pid)[0]))
    texts = [para.text for para in doc.paragraphs]
    assert "1. 题干首行" in texts                                    # Word 同样同行（且已重排）
    assert any(len(t.rows[0].cells) == 2 for t in doc.tables)        # 两图用单行表格并排


def _opts_doc(tmp_path):
    """选项含图片引用的题目（题库原始引用格式）。"""
    return [{"label": "A", "content": "![figure](asset://figures/f.webp)"},
            {"label": "B", "content": "纯文字选项"}]


async def test_option_images_migrate_and_render(client, test_db, tmp_path):
    """选项图片：创建时迁入练习资产，预览/导出均内联显示而非 Markdown 文本。"""
    from PIL import Image
    practice = await _create_practice(client, test_db, tmp_path,
                                      content="题干无图", options=_opts_doc(tmp_path))
    pid = practice["id"]
    detail = (await client.get(f"/api/practices/{pid}/detail")).json()
    opts_block = next(b for q in detail["sections"][0]["questions"] for b in q["blocks"]
                      if b["block_type"] == "options")
    opt_a = opts_block["content"][0]["content"]
    assert "asset://practice/" in opt_a and "figures/f.webp)" not in opt_a  # 已迁入资产
    assets = (await client.get(f"/api/practices/{pid}/assets-list")).json()["assets"]
    assert len(assets) == 1                                                    # 幂等：不重复复制
    (await client.get(f"/api/practices/{pid}/detail"))
    assert len((await client.get(f"/api/practices/{pid}/assets-list")).json()["assets"]) == 1

    p = await _load_with_blocks(test_db, pid)
    html = build_practice_html(p, pid)
    assert "![" not in html and "max-height:3.4em" in html                     # HTML 内联 <img>
    doc = Document(io.BytesIO(docx_export.build_docx(p, pid)[0]))
    opt_paras = [para for para in doc.paragraphs if para.text.startswith("A.")]
    assert opt_paras and any("graphic" in r._r.xml for r in opt_paras[0].runs)  # Word 行内图
    assert "![figure]" not in "".join(para.text for para in doc.paragraphs)


async def test_docx_math_omml_and_fonts(client, test_db, tmp_path):
    """Word 公式内嵌为 OMML 数学对象（行内 + 行间）；英文字体 Times New Roman、中文宋体。"""
    from docx.oxml.ns import qn
    practice = await _create_practice(client, test_db, tmp_path,
                                      content="直角三角形中满足 $a^2+b^2=c^2$，求 $c$。$$E=mc^2$$")
    pid = practice["id"]
    await client.get(f"/api/practices/{pid}/detail")
    p = await _load_with_blocks(test_db, pid)
    doc = Document(io.BytesIO(docx_export.build_docx(p, pid)[0]))
    xml = doc.element.xml
    assert "oMath" in xml                        # 公式内嵌为 OMML 对象
    assert "oMathPara" in xml                    # $$…$$ 为行间公式
    assert "$a^2+b^2=c^2$" not in xml            # 不残留原始 LaTeX 文本
    # 字体：西文 Times New Roman，中文宋体
    normal = doc.styles["Normal"]
    assert normal.font.name == "Times New Roman"
    assert normal._element.rPr.rFonts.get(qn("w:eastAsia")) == "宋体"
    # 预览 HTML 同样西文优先 TNR
    assert '"Times New Roman"' in build_practice_html(p, pid)


async def test_fit_width_capped(client, test_db, tmp_path):
    """fit 默认上限：宽图不超内容区 50%（800px@96dpi≈20cm → 封顶 8cm）。"""
    from PIL import Image
    ocr_dir = tmp_path / "ocr" / "d" / "figures"
    ocr_dir.mkdir(parents=True, exist_ok=True)
    Image.new("RGB", (800, 100)).save(ocr_dir / "big.png")
    practice = await _create_practice(client, test_db, tmp_path,
                                      content="题干\n![图](asset://figures/big.png)")
    pid = practice["id"]
    await client.get(f"/api/practices/{pid}/detail")
    p = await _load_with_blocks(test_db, pid)
    doc = Document(io.BytesIO(docx_export.build_docx(p, pid)[0]))
    shape = doc.inline_shapes[0]
    content_width = docx_export.A4_W - 2 * docx_export.Cm(2.5)   # 默认 normal 边距 25mm
    assert shape.width == int(content_width / 2)                   # 恰好封顶到 50%
